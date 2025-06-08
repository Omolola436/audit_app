import os
import json
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.chart import PieChart, Reference
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from app import app, db
from models import User, Response, Question

def generate_excel_report(user_id):
    """Generate Excel report with user responses and charts"""
    user = User.query.get(user_id)
    responses = db.session.query(Response, Question).join(Question).filter(Response.user_id == user_id).all()
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Audit Responses"
    
    # Header styling
    header_fill = PatternFill(start_color="FF8C00", end_color="FF8C00", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    # Add headers
    headers = ["Company", "Email", "Category", "Question", "Answer", "Comment", "File Upload"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # Add data
    row = 2
    for response, question in responses:
        ws.cell(row=row, column=1, value=user.company_name)
        ws.cell(row=row, column=2, value=user.email)
        ws.cell(row=row, column=3, value=question.category)
        ws.cell(row=row, column=4, value=question.question_text)
        ws.cell(row=row, column=5, value=response.answer)
        ws.cell(row=row, column=6, value=response.comment or "")
        ws.cell(row=row, column=7, value="Yes" if response.file_path else "No")
        row += 1
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{user.company_name.replace(' ', '_')}_audit_report_{timestamp}.xlsx"
    filepath = os.path.join(app.config['REPORT_FOLDER'], filename)
    wb.save(filepath)
    
    return filepath

def generate_category_chart(category_responses, category_name):
    """Generate pie chart for category responses"""
    # Count responses
    response_counts = {}
    for response in category_responses:
        answer = response.answer
        response_counts[answer] = response_counts.get(answer, 0) + 1
    
    # Create pie chart
    fig, ax = plt.subplots(figsize=(8, 6))
    colors = ['#FF8C00', '#FFA500', '#F5DEB3']  # Dark orange, light orange, beige
    
    wedges, texts, autotexts = ax.pie(response_counts.values(), 
                                     labels=response_counts.keys(),
                                     autopct='%1.1f%%',
                                     colors=colors[:len(response_counts)])
    
    ax.set_title(f'{category_name} - Response Distribution', fontsize=14, fontweight='bold')
    
    # Save chart
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    chart_filename = f"{category_name.replace(' ', '_').replace('&', 'and')}_chart_{timestamp}.png"
    chart_path = os.path.join(app.config['REPORT_FOLDER'], chart_filename)
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path

def calculate_assurance_rating(responses):
    """Calculate assurance rating based on responses"""
    if not responses:
        return "Very Limited"
    
    yes_count = sum(1 for r in responses if r.answer == "Yes")
    partial_count = sum(1 for r in responses if r.answer == "Partially")
    total = len(responses)
    
    yes_percentage = (yes_count / total) * 100
    partial_percentage = (partial_count / total) * 100
    
    if yes_percentage >= 80:
        return "High"
    elif yes_percentage >= 60 or (yes_percentage + partial_percentage) >= 80:
        return "Reasonable"
    elif yes_percentage >= 40 or (yes_percentage + partial_percentage) >= 60:
        return "Limited"
    else:
        return "Very Limited"

def generate_word_report(user_id):
    """Generate comprehensive Word audit report"""
    user = User.query.get(user_id)
    all_responses = db.session.query(Response, Question).join(Question).filter(Response.user_id == user_id).all()
    
    # Group responses by category
    categories = {}
    for response, question in all_responses:
        if question.category not in categories:
            categories[question.category] = []
        categories[question.category].append(response)
    
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{user.company_name.upper()}")
    title_run.font.size = Inches(0.25)
    title_run.bold = True
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("DATA PROTECTION AUDIT REPORT")
    subtitle_run.font.size = Inches(0.2)
    subtitle_run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B, %Y").upper())
    date_run.font.size = Inches(0.15)
    date_run.bold = True
    
    doc.add_page_break()
    
    # Executive Summary
    exec_heading = doc.add_heading('EXECUTIVE SUMMARY', level=1)
    exec_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    exec_summary = doc.add_paragraph(
        "3CONSULTING LIMITED, a licensed Data Protection Compliance Organization (DPCO), offers a comprehensive "
        "array of services encompassing various privacy disciplines. These include compliance assessments, training, "
        "program design, policy development, auditing, consulting, and other services aimed at ensuring adherence to "
        "the Nigeria Data Protection Act (NDPA) and any relevant foreign data protection laws or regulations "
        "applicable in Nigeria, as stipulated under Article 33 of the NDPA."
    )
    
    doc.add_paragraph(
        f"This audit aimed to verify the adequacy of {user.company_name}'s technical and organizational measures "
        "in guaranteeing compliance with the Nigeria Data Protection Act (NDPA). Additionally, it aimed to evaluate "
        f"the {user.company_name} monitoring mechanisms for assessing the efficacy of compliance with established "
        "policies and procedures."
    )
    
    # Scope table
    doc.add_paragraph().add_run("The audit encompassed the following scopes and descriptions:").bold = True
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SCOPE'
    hdr_cells[1].text = 'DESCRIPTION'
    
    scope_descriptions = {
        'Accountability and Governance': 'The extent to which information governance accountability, policies and procedures, performance measurement controls, and reporting mechanisms to monitor data protection compliance to NDPA is in place and in operation throughout the organization',
        'Awareness and Training': 'The provision and monitoring of staff data protection, records management and information security training, including awareness of data protection regulation requirements relating to their roles and responsibilities.',
        'Data Processing and Sharing': 'The design and operation of controls to ensure processing and sharing of personal data complies with the principles of NDPA.',
        'Administration': 'Managing/handling all aspects of data protection procedures, and practices digitally.',
        'Capturing': 'The process of collecting and acquiring data in a secure and compliant manner digitally.',
        'Actions on Security': 'Measures to safeguard data from unauthorized access, breaches, and misuse.'
    }
    
    for scope, description in scope_descriptions.items():
        row_cells = table.add_row().cells
        row_cells[0].text = scope
        row_cells[1].text = description
    
    doc.add_page_break()
    
    # Audit Summary
    summary_heading = doc.add_heading('AUDIT SUMMARY', level=1)
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create summary table
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    summary_hdr = summary_table.rows[0].cells
    summary_hdr[0].text = 'Audit Scope Area'
    summary_hdr[1].text = 'Assurance Rating'
    summary_hdr[2].text = 'Overall Opinion'
    
    # Add category ratings
    for category_name, responses in categories.items():
        rating = calculate_assurance_rating(responses)
        opinion = f"The organization demonstrates a {rating.lower()} level of assurance in {category_name.lower()}."
        
        row_cells = summary_table.add_row().cells
        row_cells[0].text = category_name
        row_cells[1].text = rating
        row_cells[2].text = opinion
    
    # General Summary
    doc.add_paragraph()
    general_heading = doc.add_heading('GENERAL SUMMARY', level=2)
    
    doc.add_paragraph(
        f"{user.company_name} has demonstrated a level of compliance with the NDPA requirements by engaging "
        "the services of the DPCO (3Consulting) prior to their data privacy audit to ensure compliance."
    )
    
    doc.add_paragraph("The audit was conducted following 3Consulting's data protection audit methodology.")
    
    doc.add_page_break()
    
    # Charts & Graphs section
    charts_heading = doc.add_heading('CHARTS & GRAPHS', level=1)
    charts_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "The Pie charts below show the summary of the percentage breakdown of the assurance rating "
        "given for each audit scope area."
    )
    
    # Generate and insert charts for each category
    for category_name, responses in categories.items():
        if responses:
            chart_path = generate_category_chart(responses, category_name)
            
            # Add category heading
            cat_heading = doc.add_heading(category_name.upper(), level=2)
            cat_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Insert chart
            try:
                doc.add_picture(chart_path, width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph(f"Chart for {category_name} could not be generated.")
            
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # Organization Details
    org_heading = doc.add_heading('ORGANIZATION DETAILS', level=1)
    org_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    org_table = doc.add_table(rows=4, cols=2)
    org_table.style = 'Table Grid'
    
    org_data = [
        ['ORGANIZATION', user.company_name],
        ['EMAIL', user.email],
        ['AUDIT DATE', datetime.now().strftime('%B %d, %Y')],
        ['AUDIT REF', f'NDPA-{user.company_name.replace(" ", "").upper()}-{datetime.now().strftime("%Y%m%d")}']
    ]
    
    for i, (label, value) in enumerate(org_data):
        org_table.rows[i].cells[0].text = label
        org_table.rows[i].cells[1].text = value
    
    # Good Practice section
    doc.add_paragraph()
    good_practice_heading = doc.add_heading('GOOD PRACTICE', level=2)
    good_practice_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        f"{user.company_name} has demonstrated commitment to compliance with the Nigeria Data Protection Act "
        "by engaging a specialized data protection compliance organization. This engagement focuses on "
        "comprehensive compliance evaluation and continuous improvement of data protection practices."
    )
    
    # Areas for Improvement
    improvement_heading = doc.add_heading('AREAS FOR IMPROVEMENT', level=2)
    improvement_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Analyze responses for recommendations
    improvement_items = []
    for category_name, responses in categories.items():
        no_responses = [r for r in responses if r.answer == "No"]
        if no_responses:
            improvement_items.append(f"Enhancement needed in {category_name} compliance measures")
    
    if not improvement_items:
        improvement_items.append("Continue maintaining current high standards of data protection compliance")
    
    for i, item in enumerate(improvement_items, 1):
        doc.add_paragraph(f"{i}. {item}")
    
    # Save document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{user.company_name.replace(' ', '_')}_NDPA_Audit_Report_{timestamp}.docx"
    filepath = os.path.join(app.config['REPORT_FOLDER'], filename)
    doc.save(filepath)
    
    return filepath
