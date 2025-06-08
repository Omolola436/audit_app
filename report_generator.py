import os
import json
import logging
from datetime import datetime
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend
import matplotlib.pyplot as plt
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from app import app
from models import User, Question, Response, Category

def generate_excel_report(user_id):
    """Generate Excel report with user responses"""
    user = User.query.get(user_id)
    responses = Response.query.filter_by(user_id=user_id).join(Question).order_by(Question.order_num).all()
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Audit Responses"
    
    # Header styling
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="FF8C00", end_color="FF8C00", fill_type="solid")
    
    # Headers
    headers = ["Company", "Email", "Category", "Question", "Answer", "Comment", "File Uploaded", "Submission Date"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # Add data
    for row, response in enumerate(responses, 2):
        ws.cell(row=row, column=1, value=user.company_name)
        ws.cell(row=row, column=2, value=user.email)
        ws.cell(row=row, column=3, value=response.question.category)
        ws.cell(row=row, column=4, value=response.question.question_text)
        ws.cell(row=row, column=5, value=response.answer or "No answer")
        ws.cell(row=row, column=6, value=response.comment or "No comment")
        ws.cell(row=row, column=7, value="Yes" if response.file_path else "No")
        ws.cell(row=row, column=8, value=response.created_at.strftime("%Y-%m-%d %H:%M:%S"))
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save file
    filename = f"{user.company_name.replace(' ', '_')}_audit_responses_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    filepath = os.path.join(app.config['REPORT_FOLDER'], filename)
    wb.save(filepath)
    
    return filepath

def generate_category_chart(category_name, responses_data, save_path):
    """Generate pie chart for a category"""
    # Count responses
    answer_counts = {}
    for response in responses_data:
        answer = response.answer or "No Answer"
        answer_counts[answer] = answer_counts.get(answer, 0) + 1
    
    if not answer_counts:
        return None
    
    # Create pie chart
    plt.figure(figsize=(8, 6))
    colors = ['#FF8C00', '#FFD700', '#F5DEB3', '#DEB887', '#D2691E']
    
    wedges, texts, autotexts = plt.pie(answer_counts.values(), 
                                       labels=answer_counts.keys(), 
                                       autopct='%1.0f%%',
                                       colors=colors[:len(answer_counts)],
                                       startangle=90)
    
    plt.title(f'{category_name.upper()}', fontsize=14, fontweight='bold', pad=20)
    
    # Style the text
    for text in texts:
        text.set_fontsize(10)
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
        autotext.set_fontsize(10)
    
    plt.axis('equal')
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return save_path

def generate_word_report(user_id):
    """Generate comprehensive Word audit report"""
    user = User.query.get(user_id)
    responses = Response.query.filter_by(user_id=user_id).join(Question).all()
    
    # Group responses by category
    categories_data = {}
    for response in responses:
        category = response.question.category
        if category not in categories_data:
            categories_data[category] = []
        categories_data[category].append(response)
    
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_heading(f'{user.company_name.upper()}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_heading('DATA PROTECTION AUDIT REPORT', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(datetime.now().strftime('%B, %Y'))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    exec_summary = doc.add_paragraph()
    exec_summary.add_run('3CONSULTING LIMITED').bold = True
    exec_summary.add_run(', a licensed Data Protection Compliance Organization (DPCO), offers a comprehensive array of services encompassing various privacy disciplines. These include compliance assessments, training, program design, policy development, auditing, consulting, and other services aimed at ensuring adherence to the Nigeria Data Protection Act (NDPA) and any relevant foreign data protection laws or regulations applicable in Nigeria, as stipulated under Article 33 of the NDPA.')
    
    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'This audit aimed to verify the adequacy of {user.company_name.upper()}\'s technical and organizational measures in guaranteeing compliance with the Nigeria Data Protection Act (NDPA). Additionally, it aimed to evaluate the {user.company_name.upper()} monitoring mechanisms for assessing the efficacy of compliance with established policies and procedures.')
    
    # Scope table
    doc.add_paragraph()
    doc.add_heading('AUDIT SCOPE', level=2)
    
    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.style = 'Table Grid'
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = scope_table.rows[0].cells
    hdr_cells[0].text = 'SCOPE'
    hdr_cells[1].text = 'DESCRIPTION'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add scope descriptions
    scopes = {
        'Governance & Accountability': 'The extent to which information governance accountability, policies and procedures, performance measurement controls, and reporting mechanisms to monitor data protection compliance to NDPA is in place and in operation throughout the organization',
        'Awareness & Training': 'The provision and monitoring of staff data protection, records management and information security training, including awareness of data protection regulation requirements relating to their roles and responsibilities.',
        'Data Processing and Sharing': 'The design and operation of controls to ensure processing and sharing of personal data complies with the principles of NDPA.',
        'Administration': 'Managing/handling all aspects of data protection procedures, and practices digitally.',
        'Capturing': 'The process of collecting and acquiring data in a secure and compliant manner digitally. This involves capturing data from various sources, such as internal systems, external databases, or third-party vendors, while adhering to established data protection policies and procedures.',
        'Actions On Security': 'Measures to safeguard data from unauthorized access, breaches, and misuse. This includes establishing access controls, authentication mechanisms, and authorization processes to ensure only authorized personnel can access and modify data.'
    }
    
    for scope, description in scopes.items():
        row_cells = scope_table.add_row().cells
        row_cells[0].text = scope
        row_cells[1].text = description
    
    # Page break
    doc.add_page_break()
    
    # Audit Summary
    doc.add_heading('AUDIT SUMMARY', level=1)
    
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'Audit Scope Area'
    hdr_cells[1].text = 'Assurance Rating'
    hdr_cells[2].text = 'Overall Opinion'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add category summaries
    for category, category_responses in categories_data.items():
        if category_responses:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = category
            
            # Calculate rating based on responses
            positive_responses = sum(1 for r in category_responses if r.answer and r.answer.lower() in ['yes', 'high', 'fully implemented'])
            total_responses = len(category_responses)
            
            if total_responses > 0:
                ratio = positive_responses / total_responses
                if ratio >= 0.8:
                    rating = "High"
                    opinion = f"The organization demonstrates a high level of assurance in {category.lower()}, maintaining strong compliance with data privacy regulations."
                elif ratio >= 0.6:
                    rating = "Reasonable"
                    opinion = f"The organization shows reasonable assurance in {category.lower()}, with some areas identified for improvement."
                else:
                    rating = "Limited"
                    opinion = f"The organization has limited assurance in {category.lower()}, requiring significant improvements to ensure compliance."
            else:
                rating = "Not Assessed"
                opinion = "No responses recorded for this category."
            
            row_cells[1].text = rating
            row_cells[2].text = opinion
    
    # Page break
    doc.add_page_break()
    
    # Charts and Graphs
    doc.add_heading('CHARTS & GRAPHS', level=1)
    doc.add_paragraph('The pie charts below show the summary of the percentage breakdown of the assurance rating given for each audit scope.')
    
    # Generate charts for each category
    chart_paths = []
    for category, category_responses in categories_data.items():
        if category_responses:
            chart_filename = f"{category.replace(' ', '_').lower()}_chart.png"
            chart_path = os.path.join(app.config['REPORT_FOLDER'], chart_filename)
            
            generated_chart = generate_category_chart(category, category_responses, chart_path)
            if generated_chart:
                chart_paths.append((category, generated_chart))
    
    # Add charts to document
    for category, chart_path in chart_paths:
        doc.add_paragraph()
        chart_para = doc.add_paragraph()
        chart_run = chart_para.add_run()
        chart_run.add_picture(chart_path, width=Inches(5))
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Page break
    doc.add_page_break()
    
    # Organization Details
    doc.add_heading('ORGANIZATION DETAILS', level=1)
    
    org_table = doc.add_table(rows=4, cols=2)
    org_table.style = 'Table Grid'
    
    org_data = [
        ['ORGANIZATION', user.company_name.upper()],
        ['EMAIL', user.email],
        ['AUDIT DATE', datetime.now().strftime('%B %d, %Y')],
        ['AUDIT REF', f'NDPA-{user.company_name.replace(" ", "").upper()[:3]}-{datetime.now().strftime("%Y%m")}']
    ]
    
    for i, (label, value) in enumerate(org_data):
        org_table.rows[i].cells[0].text = label
        org_table.rows[i].cells[1].text = value
        # Make labels bold
        for paragraph in org_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Areas for Improvement
    doc.add_paragraph()
    doc.add_heading('AREAS FOR IMPROVEMENT', level=2)
    
    improvement_para = doc.add_paragraph()
    improvement_para.add_run('Based on the audit findings, the following areas have been identified for improvement:')
    
    # Add generic recommendations based on low-scoring categories
    recommendations = []
    for category, category_responses in categories_data.items():
        if category_responses:
            positive_responses = sum(1 for r in category_responses if r.answer and r.answer.lower() in ['yes', 'high', 'fully implemented'])
            total_responses = len(category_responses)
            
            if total_responses > 0 and (positive_responses / total_responses) < 0.8:
                if category == "Governance & Accountability":
                    recommendations.append("Strengthen governance framework and accountability measures")
                elif category == "Awareness & Training":
                    recommendations.append("Enhance staff training and awareness programs")
                elif category == "Data Processing and Sharing":
                    recommendations.append("Improve data processing and sharing controls")
                elif category == "Administration":
                    recommendations.append("Enhance administrative procedures and documentation")
                elif category == "Capturing":
                    recommendations.append("Improve data collection and capturing processes")
                elif category == "Actions On Security":
                    recommendations.append("Strengthen information security measures")
    
    if not recommendations:
        recommendations.append("Continue maintaining current high standards of data protection compliance")
    
    for i, rec in enumerate(recommendations, 1):
        rec_para = doc.add_paragraph()
        rec_para.add_run(f'{i}. ').bold = True
        rec_para.add_run(rec)
    
    # Disclaimer
    doc.add_paragraph()
    doc.add_heading('DISCLAIMER', level=2)
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.add_run('This audit report is based on the information provided during the assessment period. 3Consulting Limited makes no representations or warranties regarding the completeness or accuracy of the information provided. This report is intended for the exclusive use of ').add_run(user.company_name).bold = True
    disclaimer_para.add_run(' and should not be distributed to third parties without prior written consent from 3Consulting Limited.')
    
    # Save document
    filename = f"{user.company_name.replace(' ', '_')}_audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(app.config['REPORT_FOLDER'], filename)
    doc.save(filepath)
    
    # Clean up chart files
    for _, chart_path in chart_paths:
        try:
            os.remove(chart_path)
        except:
            pass
    
    return filepath
